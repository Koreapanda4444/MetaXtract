from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Dict

from PIL import Image, TiffImagePlugin

from reportlab.pdfgen import canvas
import docx  # python-docx


def _write_jpeg_noexif(path: Path) -> None:
    img = Image.new("RGB", (64, 64), (10, 20, 30))
    img.save(path, quality=85)


def _write_jpeg_with_gps(path: Path) -> None:
    img = Image.new("RGB", (64, 64), (200, 50, 50))
    exif = img.getexif()
    gps_ifd = {
        1: "N",
        2: (
            TiffImagePlugin.IFDRational(37, 1),
            TiffImagePlugin.IFDRational(34, 1),
            TiffImagePlugin.IFDRational(0, 1),
        ),
        3: "E",
        4: (
            TiffImagePlugin.IFDRational(126, 1),
            TiffImagePlugin.IFDRational(58, 1),
            TiffImagePlugin.IFDRational(0, 1),
        ),
    }
    exif[34853] = gps_ifd
    img.save(path, exif=exif, quality=85)


def _write_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path))
    c.setTitle("MetaXtract PDF")
    c.setAuthor("MetaXtract")
    c.drawString(72, 720, "MetaXtract fixture PDF")
    c.showPage()
    c.save()


def _write_docx(path: Path) -> None:
    d = docx.Document()
    d.add_heading("MetaXtract DOCX", level=1)
    d.add_paragraph("MetaXtract fixture DOCX")
    props = d.core_properties
    props.author = "MetaXtract"
    props.title = "MetaXtract DOCX"
    d.save(str(path))


def _ffmpeg_available() -> bool:
    return shutil.which("ffmpeg") is not None


def _write_mp4(path: Path) -> bool:
    if not _ffmpeg_available():
        return False

    # Generate a tiny 1 second test pattern video (no audio).
    cmd = [
        "ffmpeg",
        "-y",
        "-f",
        "lavfi",
        "-i",
        "testsrc=size=64x64:rate=10",
        "-t",
        "1",
        "-pix_fmt",
        "yuv420p",
        str(path),
    ]
    proc = subprocess.run(cmd, check=False, capture_output=True)
    return proc.returncode == 0 and path.exists()


def ensure_fixtures(fixtures_dir: Path) -> Dict[str, bool]:
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    created: Dict[str, bool] = {}

    jpg_gps = fixtures_dir / "sample_gps.jpg"
    if not jpg_gps.exists():
        _write_jpeg_with_gps(jpg_gps)
        created[jpg_gps.name] = True
    else:
        created[jpg_gps.name] = False

    jpg_plain = fixtures_dir / "sample_noexif.jpg"
    if not jpg_plain.exists():
        _write_jpeg_noexif(jpg_plain)
        created[jpg_plain.name] = True
    else:
        created[jpg_plain.name] = False

    pdf = fixtures_dir / "sample.pdf"
    if not pdf.exists():
        _write_pdf(pdf)
        created[pdf.name] = True
    else:
        created[pdf.name] = False

    doc = fixtures_dir / "sample.docx"
    if not doc.exists():
        _write_docx(doc)
        created[doc.name] = True
    else:
        created[doc.name] = False

    mp4 = fixtures_dir / "sample.mp4"
    if not mp4.exists():
        created[mp4.name] = _write_mp4(mp4)
    else:
        created[mp4.name] = False

    return created
